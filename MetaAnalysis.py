import os
import pandas as pd
import bibtexparser
from tqdm import tqdm
import json
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from textblob import TextBlob
import pycountry
from openpyxl import Workbook
from openpyxl.styles import Alignment
import ast
from itertools import combinations
import geopandas as gpd
import pickle
import docx
import schemdraw
from schemdraw import flow
from gensim.models import Word2Vec
from sklearn.cluster import KMeans
from sklearn.decomposition import PCA
from collections import Counter
from scipy.stats import linregress
import numpy as np
from docx.shared import Cm, Inches
from docx.shared import Pt
from docx.shared import RGBColor
import copy
from collections import defaultdict
import itertools
from nltk.stem import WordNetLemmatizer
from scipy.stats import ks_2samp
from scipy.stats import ttest_ind
from scipy.stats import f_oneway, kruskal, mannwhitneyu
from statsmodels.stats.multitest import multipletests


def write_pickle_file(data, folder, file):
    """
    Writes a Python object to a pickle file.
    
    :param data: Python object to save
    :param folder: Folder path where the pickle file will be saved
    :param file: Name of the pickle file (including .pkl extension)
    """
    # Ensure the save folder exists
    os.makedirs(folder, exist_ok=True)
    
    # Construct full file path
    file_path = os.path.join(folder, file)
    
    # Write data to a pickle file
    with open(file_path, 'wb') as f:
        pickle.dump(data, f)
    print(f"Data saved successfully to: {file_path}")
    
def read_pickle_file(folder, file):
    """
    Reads a Python object from a pickle file.
    
    :param folder: Folder path where the pickle file is located
    :param file: Name of the pickle file (including .pkl extension)
    :return: Python object loaded from the pickle file
    """
    # Construct full file path
    file_path = os.path.join(folder, file)
    
    # Read data from a pickle file
    with open(file_path, 'rb') as f:
        data = pickle.load(f)
    print(f"Data loaded successfully from: {file_path}")
    return data

def process_single_bib_file(file_path, file_name):
    """
    Processes a single .bib file and extracts its entries.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as bib_file:
            bib_data = bibtexparser.load(bib_file)
            for entry in bib_data.entries:
                entry['source_file'] = file_name
            return bib_data.entries
    except Exception as e:
        print(f"Error processing file {file_name}: {e}")
        return []


def process_bib_folder_to_dataframe_multithread(folder_path, max_threads=4):
    """
    Processes all .bib files in a folder using multithreading and merges them into a DataFrame.
    """
    all_entries = []
    files = [f for f in os.listdir(folder_path) if f.endswith('.bib')]

    with ThreadPoolExecutor(max_threads) as executor:
        futures = {
            executor.submit(
                process_single_bib_file,
                os.path.join(folder_path, file_name),
                file_name
            ): file_name for file_name in files
        }
        for future in tqdm(as_completed(futures), total=len(futures), desc="Reading .bib Files:"):
            all_entries.extend(future.result())

    return pd.DataFrame(all_entries)

def get_the_first_application_year(combined_df, keyword_subset, term_to_check):
    res = {}
    for key in tqdm(keyword_subset.keys(), total=len(keyword_subset), desc = 'Process Rows:'):
        rows = keyword_subset[key]
        filtered_df = combined_df.loc[rows]
        # check the first year of the term
        first_year = filtered_df[filtered_df['title'].str.contains(term_to_check)]['year'].min()
        if not first_year:
            first_year = filtered_df[filtered_df['abstract'].str.contains(term_to_check)]['year'].min()
        if not first_year:
            first_year = filtered_df[filtered_df['keywords'].str.contains(term_to_check)]['year'].min()
        if not first_year:
            res[key] = -999
        else:
            res[key] = first_year
    return res

def read_json_as_dict(folder_path, file_name):
    """
    Reads a JSON file and saves its contents as a dictionary.
    """
    file_path = os.path.join(folder_path, file_name)
    try:
        with open(file_path, 'r', encoding='utf-8') as json_file:
            return json.load(json_file)
    except Exception as e:
        print(f"Error reading JSON file: {e}")
        return {}

def read_json_files_as_dict(json_folder, json_prefix='name_data', key_field = "name", value_field = "gender"):
    file_list = [f for f in os.listdir(json_folder) if f.startswith(json_prefix)]
    
    result_dict = {}

    for file in file_list:
        file_path = os.path.join(json_folder, file)
        if not os.path.isfile(file_path):
            print(f"File not found: {file_path}")
            continue

        try:
            with open(file_path, 'r') as json_file:
                data = json.load(json_file)

                if isinstance(data, list):  # If JSON contains a list of records
                    for record in data:
                        key = record.get(key_field)
                        value = record.get(value_field)
                        if key is not None and value is not None:
                            result_dict[key] = value
                else:
                    print(f"Unexpected JSON format in file: {file_path}")

        except Exception as e:
            print(f"Error processing file {file_path}: {e}")

    return result_dict


def country_code_to_name(country_code):
    """
    Converts a 2-letter country code to the full country name using pycountry.

    :param country_code: ISO 3166-1 alpha-2 country code (e.g., "US", "JP")
    :return: Full country name (e.g., "United States", "Japan"), or None if not found
    """
    try:
        country = pycountry.countries.get(alpha_2=country_code)
        return country.name if country else None
    except KeyError:
        return None
    
def detect_country_in_university(university_name, country_data, ignore_case=True):
    """
    Detects if a university name contains a city or region name from the specified countries.
    """
    if not university_name or not isinstance(university_name, str):
        return None

    for country, locations in country_data.items():
        if ignore_case:
            pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, locations)) + r')\b', re.IGNORECASE)
        else:
            pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, locations)) + r')\b')
        if pattern.search(university_name):
            return country
    return None

def load_cities_json_with_pycountry(file_folder, file_name):
    """
    Loads cities.json into a dictionary where the key is the full country name
    and the value is a list of cities in that country using pycountry.

    :param file_folder: Path to the folder containing the cities.json file
    :param file_name: Name of the cities.json file
    :return: Dictionary with country names as keys and city names as values
    """
    file_path = os.path.join(file_folder, file_name)

    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            city_data = json.load(file)

        # Create a dictionary with country names as keys and lists of cities as values
        cities_by_country = {}
        for entry in city_data:
            country_code = entry.get("country")
            city_name = entry.get("name")
            country_name = country_code_to_name(country_code)

            if country_name and city_name:
                if country_name not in cities_by_country:
                    cities_by_country[country_name] = []
                cities_by_country[country_name].append(city_name)

        # Sort city lists for each country
        for country_name in cities_by_country:
            cities_by_country[country_name] = sorted(cities_by_country[country_name])

        return cities_by_country

    except FileNotFoundError:
        print(f"Error: File not found at {file_path}")
        return {}
    except json.JSONDecodeError:
        print("Error: Failed to decode JSON. Ensure the file is in proper JSON format.")
        return {}

def remove_before_last_corresponding_author(input_string, keyword="(corresponding author)"):
    """
    Remove the part of the string before the last occurrence of '(corresponding author)'.

    Args:
        input_string (str): The input string.

    Returns:
        str: The modified string.
    """
    
    # Find the last occurrence of the keyword
    last_occurrence_index = input_string.rfind(keyword)
    if last_occurrence_index != -1:
        # Keep only the part after the last occurrence of the keyword
        return input_string[last_occurrence_index:]
    return input_string  # Return the original string if the keyword is not found

def process_affiliations(row, university_dict, country_data1, country_data2):
    """
    Processes a single row of affiliations to map to countries.
    """
    affiliation = row['affiliation']
    if not affiliation or not isinstance(affiliation, str):
        return []
    else:
        universities = affiliation.replace(',', ' ').split("\n")
        countries = []
        for u in universities:
            if u:
                u = remove_before_last_corresponding_author(u, keyword="(Corresponding Author)")
                u = remove_before_last_corresponding_author(u, keyword="(corresponding author)")
                # get pattern for iran
                Iran_pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, ['Iran'])) + r')\b', re.IGNORECASE)
                if re.search(Iran_pattern, u):
                    countries.append('Iran')
                else:
                    detected_country = detect_country_in_university(u, country_data1)
                    country = detected_country if detected_country else "Unknown"
                    if country == "Unknown":
                        detected_country = detect_country_in_university(u, country_data2)
                        country = detected_country if detected_country else "Unknown"
                        if country == "Unknown":
                            American_state_two_letters = ['AL', 'AK', 'AZ', 'AR', 'CA', 'CO',
                                'CT', 'DE', 'FL', 'GA', 'HI', 'ID', 'IL', 'IN', 'IA', 'KS', 'KY', 
                                'LA', 'ME', 'MD', 'MA', 'MI', 'MN', 'MS', 'MO', 'MT', 'NE', 'NV', 
                                'NH', 'NJ', 'NM', 'NY', 'NC', 'ND', 'OH', 'OK', 'OR', 'PA', 'RI', 
                                'SC', 'SD', 'TN', 'TX', 'UT', 'VT', 'VA', 'WA', 'WV', 'WI', 'WY', 'DC']
                            search_pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, American_state_two_letters)) + r')\b')
                            if re.search(search_pattern, u):
                                country = 'United States'
                            elif re.search(r'\bUSA\b', u) or re.search(r'\bU\.S\.A\.\b', u) or re.search(r'\bUS\b', u):
                                country = 'United States'
                            elif re.search(r'\bUK\b', u) or re.search(r'\bU\.K\.\b', u):
                                country = 'United Kingdom'
                            elif re.search(r'\bPR\b', u) or re.search(r'\bPuerto Rico\b', u):
                                country = 'United States'
                            else:
                                country = 'Unknown'
                        else:
                            # search CAS
                            if re.search(r'\bCAS\b', u):
                                country = 'China'
                        # if country == 'Unknown':
                        #     print(f"Unknown country in affiliation: {u}")
                    countries.append(country)
        return countries
            


def map_affiliations_to_countries(df, university_data, country_data1, country_data2, max_threads=4):
    """
    Maps affiliations in the DataFrame to countries using multithreading.
    """
    university_dict = {uni['name']: uni['country'] for uni in university_data}
    countries = {}
    for idx, row in tqdm(df.iterrows(), desc="Processing Affiliations:"):
        countries[idx] = process_affiliations(row, university_dict, country_data1, country_data2)
    # Assign the countries back to the DataFrame
    df['country'] = df.index.map(countries)
    return df

def analyze_title_sentiment_subjectivity(text):
    """
    Analyzes sentiment and subjectivity for the entire title text.
    
    :param text: String containing the title
    :return: Sentiment polarity and subjectivity of the title
    """
    if pd.isna(text) or not isinstance(text, str):
        return None, None
    blob = TextBlob(text)
    return blob.sentiment.polarity, blob.sentiment.subjectivity


def analyze_abstract_sentiment_subjectivity(text):
    """
    Analyzes sentiment and subjectivity for each sentence in the abstract text.
    
    :param text: String containing the abstract
    :return: Tuple of (average sentiment, max sentiment, min sentiment, average subjectivity)
    """
    if pd.isna(text) or not isinstance(text, str):
        return None, None, None, None

    sentences = text.split('.')
    sentiments = [TextBlob(sentence).sentiment.polarity for sentence in sentences if sentence.strip()]
    subjectivities = [TextBlob(sentence).sentiment.subjectivity for sentence in sentences if sentence.strip()]

    if not sentiments:
        return None, None, None, None

    avg_sentiment = sum(sentiments) / len(sentiments)
    max_sentiment = max(sentiments)
    min_sentiment = min(sentiments)
    avg_subjectivity = sum(subjectivities) / len(subjectivities) if subjectivities else None
    return avg_sentiment, max_sentiment, min_sentiment, avg_subjectivity


def process_sentiments_multithread(df, max_threads=4):
    """
    Performs sentiment and subjectivity analysis on title and abstract in a DataFrame using multithreading.
    
    :param df: Pandas DataFrame with 'title' and 'abstract' columns
    :param max_threads: Maximum number of threads to use
    :return: DataFrame with new columns for title and abstract sentiment/subjectivity analysis
    """
    results = {}

    def process_row(index, row):
        """
        Processes a single row for title and abstract sentiment and subjectivity analysis.
        """
        title_sentiment, title_subjectivity = analyze_title_sentiment_subjectivity(row["title"])
        abstract_sentiment = analyze_abstract_sentiment_subjectivity(row["abstract"])
        return index, title_sentiment, title_subjectivity, abstract_sentiment

    with ThreadPoolExecutor(max_threads) as executor:
        futures = {executor.submit(process_row, idx, row): idx for idx, row in df.iterrows()}
        for future in tqdm(as_completed(futures), total=len(futures), desc="Processing Sentiment and Subjectivity:"):
            idx, title_sentiment, title_subjectivity, abstract_sentiment = future.result()
            results[idx] = {
                "title_sentiment": title_sentiment,
                "title_subjectivity": title_subjectivity,
                "abstract_sentiment_avg": abstract_sentiment[0],
                "abstract_sentiment_max": abstract_sentiment[1],
                "abstract_sentiment_min": abstract_sentiment[2],
                "abstract_subjectivity_avg": abstract_sentiment[3]
            }

    # Add results back to the DataFrame in the correct order
    df['title_sentiment'] = [results[i]["title_sentiment"] for i in range(len(df))]
    df['title_subjectivity'] = [results[i]["title_subjectivity"] for i in range(len(df))]
    df['abstract_sentiment_avg'] = [results[i]["abstract_sentiment_avg"] for i in range(len(df))]
    df['abstract_sentiment_max'] = [results[i]["abstract_sentiment_max"] for i in range(len(df))]
    df['abstract_sentiment_min'] = [results[i]["abstract_sentiment_min"] for i in range(len(df))]
    df['abstract_subjectivity_avg'] = [results[i]["abstract_subjectivity_avg"] for i in range(len(df))]

    return df

def get_countries_with_admin_units():
    """
    Fetches all countries and their adjectives, and includes level-1 administrative names
    for specific countries where available.

    :return: Dictionary where keys are country names, and values are lists with:
             - Country name
             - Adjective (demonym) version of the country name
             - Level-1 administrative divisions (if available)
    """
    # List of countries to include level-1 administrative regions
    countries_with_admin_units = [
        "China", "United States", "United Kingdom", "India", "Japan",
        "Germany", "Spain", "France", "Italy", "Russia", "Korea, Republic of",
        "Indonesia", "Australia", "Thailand", "Canada", "Brazil", "Maralaysia", "Russian Federation"
    ]
    common_cases = {"United States": ["USA", "U.S.", "America", "United States of America"],
        "United Kingdom": ["UK", "Britain"],
        "Korea, Republic of": ["South Korea", "Republic of Korea", "Korea"],
        'Congo, The Democratic Republic of the': ['Democratic Republic of Congo', 'DR Congo', 'DRC', 'Congo-Kinshasa', \
            'Kinshasa', 'Zaire'],
        'Congo, Republic of the': ['Republic of Congo', 'Congo-Brazzaville', 'Brazzaville'],
        "Côte d'Ivoire": ["COTE IVOIRE", "Côte d'Ivoire", "Yamoussoukro", "COTE d'IVOIRE", "Ivory Coast"],
        "Curaçao": ["Curacao"],
        "Czechia": ["Czech Republic"],
        "Eswatini": ["Swaziland"],
        "Lao People's Democratic Republic": ["Laos"],
        "Moldova, Republic of": ["Moldova"],
        "Türkiye": ["Turkey"],
        "Central African Republic": ["CAR","Ubangi-Shari"],
        "Guinea-Bissau": ["Guinea Bissau"],
        "Cape Verde": ["Cabo Verde"],
        "Saint Kitts and Nevis": ["St. Kitts and Nevis"],
        "Saint Lucia": ["St. Lucia"],
        "North Macedonia": ["Macedonia"],
        "Myanmar": ["Burma"],
        "Palestine, State of": ["Palestine"],
        "Réunion": ["Reunion"],
        "Russian Federation": ["Russia"],
        "Sint Maarten (Dutch part)": ["Sint Maarten"],
        "Trinidad and Tobago": ["Trinidad & Tobago", "Trinidad", "Tobago"],
        'Bosnia and Herzegovina': ['Bosnia', 'Herzegovina'],
        }

    country_dict = {}

    for country in pycountry.countries:
        country_name = country.name
        if country_name == 'Congo':
            country_name = 'Congo, Republic of the'
        # Initialize the list with the country name and adjective
        country_data = [country_name]
        if country_name in common_cases:
            country_data.extend(common_cases[country_name])

        # Fetch administrative divisions for specific countries
        if country_name in countries_with_admin_units:
            admin_units = []
            try:
                subdivisions = pycountry.subdivisions.get(country_code=country.alpha_2)
                admin_units = [sub.name for sub in subdivisions]
                if country_name == 'China':
                    for i in range(len(admin_units)):
                        admin_units[i] = admin_units[i].replace('Shi', '').replace('Sheng', '').replace('Zizhiqu', '').replace('SAR', '')
                    admin_units.append('Inner Mongolia')
            except AttributeError:
                admin_units = []  # No subdivisions found
            country_data.extend(admin_units)
        # Add to the dictionary
        country_dict[country_name] = country_data
        try:
            country_common_name = country.common_name
            country_dict[country_name] = country_dict[country_name] + [country_common_name]
        except:
            continue
    if 'Western Sahara' not in country_dict:
        country_dict['Western Sahara'] = {'Western Sahara', 'Sahrawi Arab Democratic Republic', 'SADR', 'Sahrawi Republic'}
    return country_dict

def get_all_country_details_from_pycountry():
    """
    Fetches all countries from pycountry and returns a dictionary where keys are
    country names and values are lists of possible aliases (e.g., official names, common names).

    :return: Dictionary of country names and aliases
    """
    country_dict = {}
    for country in tqdm(pycountry.countries):
        aliases = [country.name]
        if hasattr(country, "official_name"):
            aliases.append(country.official_name)
        if hasattr(country, "alpha_2"):
            aliases.append(country.alpha_2)
        if hasattr(country, "alpha_3"):
            aliases.append(country.alpha_3)
        country_dict[country.name] = list(set(aliases))  # Remove duplicates
    return country_dict

def resolve_special_cases(text, countries):
    """
    Resolves special cases for country mentions.

    :param text: The text to analyze (e.g., title or abstract)
    :param countries: List of detected countries
    :return: List of resolved countries
    """
    text_lower = text.lower()

    if 'United States' in countries and 'Georgia' in countries:
        if 'county' in text_lower:
            return ['United States']
        else:
            return ['Georgia']

    if 'China' in countries and 'Mongolia' in countries:
        if 'inner mongolia' in text_lower:
            return ['China']
        else:
            return ['Mongolia']

    return countries


def detect_study_site(text, combined_country_dict):
    """
    Detects mentioned countries in the text using a combined country dictionary.
    Combines results to guess the study site.

    :param text: The text to analyze (e.g., title or abstract)
    :param combined_country_dict: A dictionary where keys are country names and values are lists of aliases
    :return: A list of countries mentioned, or 'Large-Scale'/'unknown'
    """
    if pd.isna(text) or not isinstance(text, str):
        return []

    # Detect countries
    countries_mentioned = []
    for country, aliases in combined_country_dict.items():
        # For Iran is in aliase, match with case sensitive:
        if re.search('Iran', country):
            if 'Iran' in aliases:
                pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, ['Iran'])) + r')\b', re.IGNORECASE)
                if pattern.search(text):
                    countries_mentioned.append('Iran')
        else:
            # Use word boundaries (\b) for precise matching and avoid partial matches
            pattern = re.compile(r'\b(?:' + '|'.join(map(re.escape, aliases)) + r')\b', re.IGNORECASE)
            if pattern.search(text):
                countries_mentioned.append(country)

    # Remove duplicates and sort the list
    countries_mentioned = sorted(set(countries_mentioned))

    # Resolve special cases
    countries_mentioned = resolve_special_cases(text, countries_mentioned)

    # Determine study site category
    if countries_mentioned:
        return countries_mentioned
    elif any(term in text.lower() for term in ['global', 'africa', 'asia', 'europe', 'north america', 'mideast',
                                               'latin america', 'south america', 'international', 'worldwide']):
        return ['Large-Scale']
    else:
        return ['unknown']


def process_study_site_row(index, row, combined_country_dict):
    """
    Processes a single row to determine the study site based on title and abstract.

    :param index: Row index
    :param row: A single row from the DataFrame
    :param combined_country_dict: A dictionary for country detection
    :return: Tuple (index, study_site)
    """
    title_site = detect_study_site(row['title'], combined_country_dict)
    abstract_site = detect_study_site(row['abstract'], combined_country_dict)

    # Use title information if available, otherwise use abstract
    if title_site and title_site != ['unknown']:
        study_site = title_site
    else:
        study_site = abstract_site

    return index, study_site


def analyze_study_sites_multithread(df, combined_country_dict, max_threads=4):
    """
    Analyzes the study site based on the title and abstract using multi-threading.
    Adds a new column 'study_site' to the DataFrame.

    :param df: The DataFrame containing 'title' and 'abstract' columns
    :param country_dict: A custom country dictionary
    :param max_threads: Maximum number of threads to use
    :return: The updated DataFrame with a new 'study_site' column
    """
    # Combine custom country dictionary with pycountry data
    # pycountry_dict = get_all_countries_from_pycountry()
    # combined_country_dict = {**pycountry_dict, **country_dict}

    results = {}

    with ThreadPoolExecutor(max_threads) as executor:
        futures = {executor.submit(process_study_site_row, idx, row, combined_country_dict): idx for idx, row in df.iterrows()}
        for future in tqdm(as_completed(futures), total=len(futures), desc="Analyzing Study Sites:"):
            idx, result = future.result()
            results[idx] = result

    # Add results to the DataFrame in the correct order
    df['study_site'] = [results[i] for i in range(len(df))]
    return df

def save_dataframe_to_excel(df, save_folder, save_name):
    """
    Saves a Pandas DataFrame to an Excel file in the specified folder with the specified file name.
    
    :param df: Pandas DataFrame to save
    :param save_folder: Path to the folder where the Excel file will be saved
    :param save_name: Name of the Excel file (including .xlsx extension)
    """
    # Ensure the save folder exists
    os.makedirs(save_folder, exist_ok=True)
    
    # Construct full file path
    file_path = os.path.join(save_folder, save_name)
    
    # Save DataFrame to Excel
    try:
        df.to_excel(file_path, index=False, engine='openpyxl')  # Use openpyxl for better compatibility
        print(f"DataFrame saved successfully to: {file_path}")
    except Exception as e:
        print(f"Failed to save DataFrame to Excel: {e}")
        
def save_dataframe_to_excel_multiline(df, save_folder, save_name):
    """
    Saves a Pandas DataFrame to an Excel file in the specified folder with the specified file name,
    ensuring multi-line cells are preserved with text wrapping.
    
    :param df: Pandas DataFrame to save
    :param save_folder: Path to the folder where the Excel file will be saved
    :param save_name: Name of the Excel file (including .xlsx extension)
    """
    # Ensure the save folder exists
    os.makedirs(save_folder, exist_ok=True)
    
    # Construct full file path
    file_path = os.path.join(save_folder, save_name)
    
    try:
        # Save DataFrame to Excel using openpyxl
        with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Sheet1')

            # Access the workbook and worksheet
            worksheet = writer.sheets['Sheet1']

            # Apply text wrapping to all cells
            for row in tqdm(worksheet.iter_rows(), desc="Writing Dataframe into Excel File:"):
                for cell in row:
                    cell.alignment = Alignment(wrap_text=True)

        print(f"DataFrame saved successfully to: {file_path}")
    except Exception as e:
        print(f"Failed to save DataFrame to Excel: {e}")
        
def load_excel_to_dataframe(file_folder, file_name):
    """
    Loads an Excel file into a Pandas DataFrame.
    
    :param file_path: Path to the Excel file
    :return: Pandas DataFrame containing the data from the Excel file
    """
    file_path = os.path.join(file_folder, file_name)
    try:
        # Load the Excel file into a DataFrame
        df = pd.read_excel(file_path, engine='openpyxl')
        print(f"Excel file loaded successfully from: {file_path}")
        return df
    except Exception as e:
        print(f"Failed to load Excel file: {e}")
        return None
    
def read_excel_files_by_keywords(folder_path, keyword1, keyword2):
    """
    Reads all Excel files in a folder that match the specified keywords or years in their filenames
    and combines them into a single DataFrame.
    
    :param folder_path: Path to the folder containing the Excel files
    :param keyword1: First keyword to search for in filenames (str, list, or 'dddd' for year)
    :param keyword2: Second keyword to search for in filenames (str, list, or 'dddd' for year)
    :return: A Pandas DataFrame containing the combined data
    """
    # Convert keywords to lists if they are not already
    if isinstance(keyword1, str) and keyword1 == 'dddd':
        keyword1 = r"\b\d{4}\b"  # Regex for a 4-digit year
    elif isinstance(keyword1, str) or isinstance(keyword1, int):
        keyword1 = [str(keyword1)]
    
    if isinstance(keyword2, str) and keyword2 == 'dddd':
        keyword2 = r"\b\d{4}\b"  # Regex for a 4-digit year
    elif isinstance(keyword2, str) or isinstance(keyword2, int):
        keyword2 = [str(keyword2)]

    combined_df = pd.DataFrame()

    # Iterate through files in the folder
    for file_name in os.listdir(folder_path):
        # Check if the file is an Excel file
        if file_name.endswith(".xlsx"):
            # Check for year pattern or keywords in the filename
            keyword1_match = (
                re.search(keyword1, file_name) if isinstance(keyword1, str) else any(k1 in file_name for k1 in keyword1)
            )
            keyword2_match = (
                re.search(keyword2, file_name) if isinstance(keyword2, str) else any(k2 in file_name for k2 in keyword2)
            )
            
            # If both keywords match, process the file
            if keyword1_match and keyword2_match:
                file_path = os.path.join(folder_path, file_name)
                try:
                    # Read the Excel file and append to the combined DataFrame
                    df = pd.read_excel(file_path, engine='openpyxl')
                    combined_df = pd.concat([combined_df, df], ignore_index=True)
                    print(f"Loaded: {file_name}")
                except Exception as e:
                    print(f"Failed to read {file_name}: {e}")

    return combined_df

def remove_duplicates_by_unique_id(df, unique_id_column):
    """
    Removes potential duplicate rows from the DataFrame based on a unique ID column.
    
    :param df: Pandas DataFrame to process
    :param unique_id_column: Name of the column containing unique IDs
    :return: DataFrame with duplicates removed
    """
    if unique_id_column not in df.columns:
        raise ValueError(f"Column '{unique_id_column}' not found in the DataFrame.")
    
    # Drop duplicate rows based on the unique ID column
    df = df.drop_duplicates(subset=[unique_id_column], keep="first").reset_index(drop=True)
    print(f"Removed duplicates. Remaining rows: {len(df)}")
    return df

def analyze_yearly_publications(df, year_column, country_column, countries_of_interest=None):
    """
    Analyzes the DataFrame to extract yearly publication statistics, focusing on total 
    publications and specific countries. Subsets data by year and processes country column row by row.

    :param df: Pandas DataFrame containing the data
    :param year_column: Name of the column containing year information
    :param country_column: Name of the column containing country information (string or list)
    :param countries_of_interest: List of specific countries to report (optional)
    :return: A dictionary with yearly publication statistics
    """
    # Initialize the result dictionary
    yearly_report = {}

    # Ensure year column is treated as string for consistent grouping
    df[year_column] = df[year_column].astype(str)

    # Ensure countries_of_interest is case-insensitive
    if countries_of_interest:
        countries_of_interest = [country.strip().lower() for country in countries_of_interest]

    # Group data by year
    for year, year_group in tqdm(df.groupby(year_column), desc="Processing grouping by year and countries"):
        # Initialize yearly data
        yearly_data = {
            "total_publications": len(year_group),
            "countries": {}
        }

        # If countries of interest are provided, calculate their statistics
        if countries_of_interest:
            country_counts = {country: 0 for country in countries_of_interest}

            # Iterate over each row in the year subset
            for _, row in year_group.iterrows():
                # Get the value of the country column
                temp_countries = row[country_column]
                if isinstance(temp_countries, list):
                    countries_in_row = ast.literal_eval(temp_countries[0])
                else:
                    countries_in_row = ast.literal_eval(temp_countries)
                             
                # Count occurrences of each country of interest
                for country in countries_of_interest:
                    if any(c.lower() == country for c in countries_in_row):
                        country_counts[country] += 1

            # Calculate percentages and add to yearly data
            for country, count in country_counts.items():
                yearly_data["countries"][country] = {
                    "count": count,
                    "percentage": (count / len(year_group)) * 100 if len(year_group) > 0 else 0
                }

        # Add the yearly data to the report
        yearly_report[year] = yearly_data

    return yearly_report

def create_keyword_subset_multithread(df, keyword_dict, title_column="title", abstract_column="abstract", max_threads=4):
    """
    Creates a dictionary where each keyword maps to a list of row indices 
    in the DataFrame where the keyword or its associated items are mentioned 
    in the title or abstract. Uses multi-threading to accelerate the process.

    :param df: The DataFrame containing the data.
    :param keyword_dict: A dictionary where keys are keywords and values are lists of associated terms.
    :param title_column: The column name for titles (default: "title").
    :param abstract_column: The column name for abstracts (default: "abstract").
    :param max_threads: Maximum number of threads to use (default: 4).
    :return: A dictionary where each keyword maps to a list of subset indices.
    """
    # Initialize the keyword subset dictionary
    keyword_subset = {keyword: [] for keyword in keyword_dict}

    def process_row(idx, row):
        """
        Processes a single row to check if it matches any keywords.
        """
        matches = {keyword: [] for keyword in keyword_dict}
        text = f"{row[title_column]} {row[abstract_column]}".lower()

        for keyword, terms in keyword_dict.items():
            if any(term.lower() in text for term in terms):
                matches[keyword].append(idx)

        return matches

    # Multi-threaded execution
    with ThreadPoolExecutor(max_threads) as executor:
        futures = {executor.submit(process_row, idx, row): idx for idx, row in df.iterrows()}
        for future in tqdm(as_completed(futures), total=len(futures), desc="Processing rows"):
            result = future.result()
            for keyword, indices in result.items():
                keyword_subset[keyword].extend(indices)

    return keyword_subset

def create_keyword_subset(df, keyword_dict, title_column="title", abstract_column="abstract"):
    """
    Creates a dictionary where each keyword maps to a list of row indices 
    in the DataFrame where the keyword or its associated items are mentioned 
    in the title or abstract. Uses tqdm to show the process.

    Exact matching (with word boundaries) is applied if the keyword/term is 3 words or fewer, 
    while substring matching is applied for terms with more than 3 words.

    :param df: The DataFrame containing the data.
    :param keyword_dict: A dictionary where keys are keywords and values are lists of associated terms.
    :param title_column: The column name for titles (default: "title").
    :param abstract_column: The column name for abstracts (default: "abstract").
    :return: A dictionary where each keyword maps to a list of subset indices.
    """
    # Initialize the keyword subset dictionary
    keyword_subset = {keyword: [] for keyword in keyword_dict}

    # Iterate over each row in the DataFrame with tqdm for progress
    for idx, row in tqdm(df.iterrows(), total=len(df), desc="Processing rows"):
        # Combine title and abstract for search
        text = f"{row[title_column]} {row[abstract_column]}".lower() if pd.notnull(row[title_column]) and pd.notnull(row[abstract_column]) else ""

        # Check for each keyword and its associated items
        for keyword, terms in keyword_dict.items():
            for term in terms:
                term_lower = term.lower()
                if len(term) <= 4:
                    # Exact matching using word boundaries
                    
                    
                    # Tokenize the text into words for exact matching
                    words_in_text = set(re.findall(r'\b\w+\b', text))  # Extracts full words using regex                 
                    
                    # if re.search(rf'\b{re.escape(term_lower)}\b', text):
                    #     keyword_subset[keyword].append(idx)
                    #     break  # No need to check further terms for this keyword
                    if term_lower in words_in_text:
                        keyword_subset[keyword].append(idx)
                        break  # No need to check further terms for this keyword
                else:
                    # Substring matching
                    if term_lower in text:
                        keyword_subset[keyword].append(idx)
                        break  # No need to check further terms for this keyword

    return keyword_subset

def filter_df_based_on_year_range(df, year_column, start_year, end_year):
    """
    Filters the DataFrame to include only rows within a specified range of years.

    :param df: Pandas DataFrame containing the data
    :param year_column: Name of the column containing year information
    :param start_year: The start year of the range (inclusive)
    :param end_year: The end year of the range (inclusive)
    :return: A filtered DataFrame containing only the specified years
    """
    return df[(df[year_column] >= start_year) & (df[year_column] <= end_year)]


def filter_keyword_subset_by_year(df, keyword_subset, year_column, start_year, end_year):
    """
    Filters the keyword subset to include only entries within a specified year range.

    :param df: DataFrame containing the year information.
    :param keyword_subset: Dictionary where keywords map to lists of row indices.
    :param year_column: Column name in the DataFrame containing year data.
    :param start_year: Start year as a string (format: 'yyyy').
    :param end_year: End year as a string (format: 'yyyy').
    :return: A filtered keyword subset dictionary.
    """

    # Filter the DataFrame to get the indices in the year range
    valid_indices = set(df[(df[year_column] >= start_year) & (df[year_column] <= end_year)].index)

    # Filter the keyword subset
    filtered_keyword_subset = {
        keyword: [idx for idx in indices if idx in valid_indices]
        for keyword, indices in keyword_subset.items()
    }

    return filtered_keyword_subset

def remove_unknown(lst):
    """
    Removes all occurrences of 'unknown' from a list.

    :param lst: List of elements
    :return: List with 'unknown' removed
    """
    return [item for item in lst if item != "unknown"]

def get_country_code(country_name):
    """
    Converts a country name to its two-letter code.
    If the name is already a code, validates it.
    """
    try:
        if country_name == 'unknown' or country_name == 'Unknown' or country_name == 'Large-Scale':
            return None
        # Check if input is already a two-letter or three-letter code
        if len(country_name) == 2 and pycountry.countries.get(alpha_2=country_name.upper()):
            res = country_name.upper()
        elif len(country_name) == 3 and pycountry.countries.get(alpha_3=country_name.upper()):
            res = pycountry.countries.get(alpha_3=country_name.upper()).alpha_2
        # Convert from full name
        else:
            if country_name == 'Russia':
                res = 'RU'
            elif country_name == 'Congo, Republic of the':
                res = 'CG'
            elif country_name == 'Iran':
                res = 'IR'
            elif country_name == 'Democratic Republic of Congo':
                res = 'CD'
            elif country_name == 'Republic of Congo':
                res = 'CG'
            elif country_name == "Cote d'Ivoire":
                res = 'CI'
            elif country_name == 'Brunei':
                res = 'BN'
            elif country_name == 'Cape Verde':
                res = 'CV'
            elif country_name == 'East Timor':
                res = 'TL'
            elif country_name == 'Micronesia (country)' or country_name == 'Micronesia':
                res = 'FM'
            elif country_name == 'Palestine':
                res = 'PS'
            elif country_name == 'Turkey':
                res = 'TR'
            elif country_name == 'United States Virgin Islands':
                res = 'VI'
            elif country_name == 'Urkanie' or country_name == 'Ukraine':
                res = 'UA'
            elif country_name == 'Czech':
                res = 'CZ'
            elif country_name == 'Netherlands Antilles':
                res = 'AN'
            elif country_name == 'Kosovo':
                res = 'RS'
            elif country_name == 'Vatican City':
                res = 'VA'
            else:
                res = pycountry.countries.get(name=country_name).alpha_2
        if res == 'HK' or res == 'TW':
            res = 'CN'
        return res
    except:
        
        try:
            res = pycountry.countries.get(common_name=country_name).alpha_2
            
            return res
        except:
            #print(country_name) 
            return None  # If no valid match is found

def analyze_country_pairs(df, country_column):
    """
    Analyzes a DataFrame to count pairs of countries.

    :param df: DataFrame containing country data
    :param country_column: Column name containing country lists
    :return: Dictionary of country pairs and their counts
    """
    country_pair_dict = {}

    # Iterate through rows with tqdm for progress tracking
    for _, row in tqdm(df.iterrows(), total=len(df), desc="Processing rows"):
        # Extract and clean country list
        countries = row[country_column]
        if isinstance(countries, list):
            countries = ast.literal_eval(countries[0])
        else:
            countries = ast.literal_eval(countries)
        
        countries = list(set(remove_unknown(countries)))  # Remove duplicates and "unknown"

        # Convert all countries to two-letter codes
        country_codes = [get_country_code(country) for country in countries]
        country_codes = list(set(country_codes))  # Remove duplicates
        country_codes = [code for code in country_codes if code]  # Filter out invalid codes

        # Generate all unique pairs (combinations)
        if len(country_codes) > 1:
            for pair in combinations(sorted(country_codes), 2):
                if pair in country_pair_dict:
                    country_pair_dict[pair] += 1
                else:
                    country_pair_dict[pair] = 1

    return country_pair_dict

def analyze_country_counts(df, country_column, is_remove_duplicate = True):
    """
    Analyzes a DataFrame to count pairs of countries.

    :param df: DataFrame containing country data
    :param country_column: Column name containing country lists
    :return: Dictionary of country pairs and their counts
    """
    country_count_dict = {}

    # Iterate through rows with tqdm for progress tracking
    for _, row in tqdm(df.iterrows(), total=len(df), desc="Processing rows"):
        # Extract and clean country list
        countries = row[country_column]
        if isinstance(countries, list):
            countries = ast.literal_eval(countries[0])
        else:
            countries = ast.literal_eval(countries)

        if is_remove_duplicate:
            countries = list(set(remove_unknown(countries)))  # Remove duplicates and "unknown"
        else:
            countries = remove_unknown(countries)
        # Convert all countries to two-letter codes
        country_codes = [get_country_code(country) for country in countries]
        country_codes = [code for code in country_codes if code]  # Filter out invalid codes
        if country_codes:
            for code in country_codes:
                if code in country_count_dict:
                    country_count_dict[code] += 1
                else:
                    country_count_dict[code] = 1

    return country_count_dict

def extract_disputed_regions_by_fid(shp_folder, shp_file, target_fids):
    """
    Extract geometries for disputed regions using FIDs from a shapefile.

    Parameters:
    - shp_folder: str, path to the folder containing shapefiles.
    - shp_file: str, name of the shapefile to read.
    - target_fids: list of int, FIDs of disputed regions to extract.

    Returns:
    - dict: A dictionary with FIDs as keys and geometries as values.
    """
    # Load the shapefile
    shp_path = shp_folder + shp_file
    disputed_regions = gpd.read_file(shp_path)

    # Extract specified FIDs
    region_geometries = {}
    for fid in target_fids:
        if fid in disputed_regions.index:
            region_geometries[fid] = disputed_regions.loc[fid].geometry
        else:
            print(f"FID {fid} not found in the shapefile.")

    return region_geometries

def create_study_site_subset(df, study_site_column="study_site"):
    """
    Creates a dictionary where each unique study site (country) maps to a list of row indices 
    in the DataFrame where the study site is mentioned.

    Parameters:
    - df: The DataFrame containing the data.
    - study_site_column: The column name containing the study site information.

    Returns:
    - study_site_subset: A dictionary where keys are unique study sites and values are lists of row indices.
    """
    # Initialize the dictionary to store subsets
    study_site_subset = {}

    # Iterate over the DataFrame to process the study site column
    for idx, row in df.iterrows():
        study_sites = row[study_site_column]

        # Ensure the study site data is a list of countries (e.g., already processed correctly)
        if isinstance(study_sites, list):
            for site in study_sites:
                if site not in study_site_subset:
                    study_site_subset[site] = []
                study_site_subset[site].append(idx)
        elif isinstance(study_sites, str):  # Handle case if it's a string (e.g., '["Country1", "Country2"]')
            try:
                study_sites_list = ast.literal_eval(study_sites)  # Convert string to list
                for site in study_sites_list:
                    if site not in study_site_subset:
                        study_site_subset[site] = []
                    study_site_subset[site].append(idx)
            except Exception as e:
                print(f"Error processing row {idx} with study_site={study_sites}: {e}")

    return study_site_subset

def get_subset_of_keys(subsets, keys):
    res = {}
    for key in keys:
        if key in subsets:
            res[key] = subsets[key]
    return res
def process_keyword_pairs(df, keyword_dict, title_column="title", abstract_column="abstract", min_word_len=4):
    """
    Processes a DataFrame to count keyword pairs and individual keyword counts for each category.

    Parameters:
    - df: DataFrame containing the data.
    - keyword_dict: Dictionary where keys are categories and values are lists of associated keywords.
    - title_column: Column name for titles (default: "title").
    - abstract_column: Column name for abstracts (default: "abstract").
    - min_word_len: Minimum length for substring matching; words <= this length must match exactly.

    Returns:
    - pair_count: Dictionary where keys are (keyword1, keyword2) tuples and values are counts.
    - keyword_info: Dictionary where keys are keywords and values are dictionaries with count and category info.
    """
    from collections import defaultdict
    import itertools
    import re

    pair_count = defaultdict(int)
    keyword_info = defaultdict(lambda: {"count": 0, "category": None})

    # Flatten all keywords into a single list with their categories
    all_keywords = []
    for category, keywords in keyword_dict.items():
        all_keywords.extend([(kw.lower(), category) for kw in keywords])

    # Extract just the keywords for pair counting
    keyword_list = [kw[0] for kw in all_keywords]

    for _, row in tqdm(df.iterrows(), total=len(df)):
        # Combine title and abstract into one lowercase text for searching
        text = f"{row[title_column]} {row[abstract_column]}".lower()

        # Tokenize the text into words for exact matching
        words_in_text = set(re.findall(r'\b\w+\b', text))  # Extracts full words using regex

        # Find keywords in the text based on length rule
        found_keywords = set()
        for keyword in keyword_list:
            if len(keyword) <= min_word_len:  # Exact match required for short words
                if keyword in words_in_text:
                    found_keywords.add(keyword)
            else:  # Allow substring matching for longer words
                if keyword in text:
                    found_keywords.add(keyword)

        # Count individual keywords and assign category
        for keyword in found_keywords:
            category = next(cat for kw, cat in all_keywords if kw == keyword)
            keyword_info[keyword]["count"] += 1
            keyword_info[keyword]["category"] = category

        # Count keyword pairs
        for keyword1, keyword2 in itertools.combinations(sorted(found_keywords), 2):
            pair_count[(keyword1, keyword2)] += 1

    return dict(pair_count), dict(keyword_info)

def process_two_topic_pairs(df, topic1, topic2, title_column="title", abstract_column="abstract", min_word_len=4):
    """
    Processes a DataFrame to count keyword pairs and individual keyword counts for each category.

    Parameters:
    - df: DataFrame containing the data.
    - topic1: Dictionary where keys are categories and values are lists of associated keywords (Topic Group 1).
    - topic2: Dictionary where keys are categories and values are lists of associated keywords (Topic Group 2).
    - title_column: Column name for titles (default: "title").
    - abstract_column: Column name for abstracts (default: "abstract").
    - min_word_len: Minimum length for substring matching; words <= this length must match exactly.

    Returns:
    - pair_count: Dictionary where keys are (keyword1, keyword2) tuples and values are counts.
    - keyword_info: Dictionary where keys are keywords and values are dictionaries with count and category info.
    """

    pair_count = defaultdict(int)  # Stores counts for (topic1_key, topic2_key) co-occurrences
    keyword_info = defaultdict(lambda: {"count": 0, "category": None})  # Stores individual counts

    # Standardize topics
    topic_dict_1, topic_dict_2 = standardize_topics(topic1), standardize_topics(topic2)

    for _, row in tqdm(df.iterrows(), total=len(df)):
        # Combine title and abstract into one lowercase text for searching
        text = f"{row[title_column]} {row[abstract_column]}".lower()

        # Track detected keywords from each group
        detected_topic1 = set()
        detected_topic2 = set()

        # Tokenize the text into words for exact matching
        words_in_text = set(re.findall(r'\b\w+\b', text))  # Extracts full words using regex

        # Search for topic1 keywords in the text
        for key, word_list in topic_dict_1.items():
            for word in word_list:
                if len(word) <= min_word_len:  # Exact match required for short words
                    if word in words_in_text:
                        keyword_info[key]["count"] += 1
                        keyword_info[key]["category"] = "topics1"
                        detected_topic1.add(key)
                        break  # No need to check further if we already found a match
                else:  # Allow substring matching for longer words
                    if word in text:
                        keyword_info[key]["count"] += 1
                        keyword_info[key]["category"] = "topics1"
                        detected_topic1.add(key)
                        break

        # Search for topic2 keywords in the text
        for key, word_list in topic_dict_2.items():
            for word in word_list:
                if len(word) <= min_word_len:  # Exact match required for short words
                    if word in words_in_text:
                        keyword_info[key]["count"] += 1
                        keyword_info[key]["category"] = "topics2"
                        detected_topic2.add(key)
                        break
                else:  # Allow substring matching for longer words
                    if word in text:
                        keyword_info[key]["count"] += 1
                        keyword_info[key]["category"] = "topics2"
                        detected_topic2.add(key)
                        break

        # Update pair counts if both topic1 and topic2 keywords are present
        for key1, key2 in itertools.product(detected_topic1, detected_topic2):
            pair_count[(key1, key2)] += 1

    return pair_count, keyword_info


def standardize_topics(topics):
    if isinstance(topics, dict):
        res = copy.deepcopy(topics)
        for k, v in topics.items():
            res[k] = [kw.lower() for kw in v] + [str(k).lower()]
    elif isinstance(topics, list):
        res = {}
        for k in topics:
            res[k] = [str(k).lower()]
    else:
        res = None
    return res
        
def normalize_key(key):
    """
    Normalize the key by:
    - Removing spaces, hyphens, and underscores (to merge 'heavy metal', 'heavy-metal', 'heavy_metal')
    - Converting to lowercase (initially)
    - Lemmatizing to get the base word form (e.g., 'running' -> 'run', 'cats' -> 'cat')
    """
    lemmatizer = WordNetLemmatizer()
    words = re.split(r'[_\-\s&]+', key)  # Split words on separators
    words = [lemmatizer.lemmatize(w.lower()) for w in words if w]  # Lemmatize each word
    final_word = ''.join(words)  # Convert to base form
    final_word = lemmatizer.lemmatize(lemmatizer.lemmatize(final_word, pos='v'), pos='n')
    #words = [lemmatizer.lemmatize(lemmatizer.lemmatize(w, pos='v'), pos='n') for w in words]  # Lemmatize again to handle verb forms
    
    return final_word  # Convert to base form

def merge_duplicate_keys(original_dict):
    """
    Merges duplicate keys in a dictionary based on plural forms, case sensitivity,
    word separation, hyphens, underscores, and different word forms.

    Parameters:
    - original_dict: Dictionary with possibly duplicate keys.

    Returns:
    - merged_dict: Dictionary with merged keys.
    """
    merged_dict = defaultdict(list)
    key_mapping = defaultdict(set)

    # Normalize and track original keys
    for key in original_dict:
        norm_key = normalize_key(key)
        key_mapping[norm_key].add(key)  # Store all variations of the same normalized key

    # Merge values and select the appropriate key
    for norm_key, key_variations in key_mapping.items():
        merged_values = []
        for key in key_variations:
            merged_values.extend(original_dict[key])  # Merge values
        
        # **Choose the key following the priority rules:**
        # - If a fully uppercase version exists, use that (e.g., "HEAVY_METALS")
        # - Otherwise, use the lowercase version (e.g., "heavy metal")
        final_key = next((k for k in key_variations if k.isupper()), min(key_variations, key=str.lower))
        merged_dict[final_key] = list(set(merged_values))  # Keep unique values

    return dict(merged_dict)

def generate_subset_stats(df, keyword_subset, processing_column):
    subset_dict = {}
    for key in keyword_subset:
        subset_dict[key] = df.loc[keyword_subset[key]][processing_column].describe().to_dict()
    return subset_dict

def generate_yearly_stats(df, keyword_subset, year_column='year'):
    df[year_column] = df[year_column].fillna(0)
    df[year_column] = df[year_column].astype(int) 

    # Initialize a dictionary to store yearly counts for each keyword
    unique_years = sorted(df[year_column].unique())
    yearly_counts = {keyword: {year: 0 for year in unique_years} for keyword in keyword_subset}

    # Count yearly occurrences for each keyword
    for keyword, indices in keyword_subset.items():
        subset = df.loc[indices]
        for year, count in subset[year_column].value_counts().items():
            yearly_counts[keyword][year] += count
    return yearly_counts

def create_VOSViewer_json_files(keyword_count, keyword_pairs, categories, save_folder, save_name="VOSViewer.json", is_recreate_clusters = False):
    os.makedirs(save_folder, exist_ok=True)
    item_list, link_list = [], []
    count = 1
    if is_recreate_clusters:
        cluster_ids = {}
        for i in categories.keys():
            cluster_ids[i] = count
            count += 1
    else:
        cluster_ids = categories
    
    item_id = 1
    item_ids = {}
    for j in keyword_count.keys():
        temp_name = str(j)
        temp_cat = keyword_count[j].get('category')
        temp_count = keyword_count[j].get('count')
    
        cluster_id = cluster_ids[temp_cat]
        item_list.append({"id": item_id, "label": temp_name, "cluster": cluster_id, "weights": {"count":temp_count}})
        item_ids[temp_name] = item_id
        item_id += 1
    for k in keyword_pairs.keys():
        temp_source = str(k[0])
        temp_target = str(k[1])
        temp_weight = keyword_pairs[k]
        source_id = item_ids[temp_source]
        target_id = item_ids[temp_target]
        link_list.append({"source_id": source_id, "target_id": target_id, "strength": temp_weight})
    for k1 in keyword_count.keys():
        for k2 in keyword_count.keys():
            if k1 == k2:
                continue
            else:
                temp_source = str(k1)
                temp_target = str(k2)
                if (temp_source, temp_target) not in keyword_pairs.keys() and (temp_target, temp_source) not in keyword_pairs.keys():
                    temp_weight = 0
                    source_id = item_ids[temp_source]
                    target_id = item_ids[temp_target]
                    link_list.append({"source_id": source_id, "target_id": target_id, "strength": temp_weight})
    json_dict = {"network":{"items": item_list, "links": link_list}}
    # save the json file
    with open(f'{save_folder}{save_name}', 'w') as f:
        f.write(str(json_dict))

def analyze_journal_publications(df, journal_column="journal", year_column="year", subset_list=None):
    """
    Analyze the top 10 journals by publication count and their yearly publication counts.

    Parameters:
    - df: DataFrame containing the journal and year data.
    - journal_column: Column name for journals.
    - year_column: Column name for publication years.
    - subset_list: List of indices to include. If None or empty, include all rows.

    Returns:
    - top_journals: DataFrame of top 10 journals with publication counts.
    - yearly_counts: DataFrame of yearly publication counts for top 10 journals.
    """
    # Subset the DataFrame if subset_list is provided
    if subset_list:
        df = df.loc[subset_list]

    # Count the top 10 journals
    journal_counts = df[journal_column].value_counts().head(10)
    top_journals = journal_counts.reset_index()
    top_journals.columns = [journal_column, "publication_count"]

    # Filter the DataFrame for only top 10 journals
    top_journals_list = top_journals[journal_column].tolist()
    filtered_df = df[df[journal_column].isin(top_journals_list)]

    # Group by year and journal to count yearly publications
    yearly_counts = (
        filtered_df.groupby([year_column, journal_column])
        .size()
        .reset_index(name="yearly_publications")
    )

    # Pivot the yearly counts for better visualization
    yearly_counts_pivot = yearly_counts.pivot(
        index=year_column, columns=journal_column, values="yearly_publications"
    ).fillna(0)

    return top_journals, yearly_counts_pivot

def calculate_average_impact_factor(df, keyword_subset, impact_factor_col="impact_factor"):
    """
    Calculate the average impact factor for each keyword subset.

    Parameters:
    - df: DataFrame containing the data.
    - keyword_subset: Dictionary with keywords as keys and row indices as values.
    - impact_factor_col: Column name for impact factor values.

    Returns:
    - A dictionary with keywords as keys and their average impact factor as values.
    """
    keyword_impact_factors = {}

    for keyword, indices in keyword_subset.items():
        # Get the subset of the DataFrame
        subset_df = df.loc[indices]

        # Calculate the average impact factor, ignoring missing values
        if not subset_df.empty:
            average_if = subset_df[impact_factor_col].mean(skipna=True)
        else:
            average_if = None  # Handle cases with no data

        keyword_impact_factors[keyword] = average_if

    return keyword_impact_factors

def get_country_region(country_name):
    '''
    Get the region of a country based on its name.
    country name could be a name, 2-letter code, or 3-letter code.
    return the region of the country.
    '''
    continent_mapping = {
        "Asia": ["AF", "AM", "AZ", "BH", "BD", "BT", "BN", "KH", "CN", "CY", "GE", "IN", "ID", "IR", "IQ", "IL", "JP",
                 "JO", "KZ", "KW", "KG", "LA", "LB", "MY", "MV", "MN", "MM", "NP", "KP", "OM", "PK", "PS", "PH", "QA",
                 "SA", "SG", "KR", "LK", "SY", "TJ", "TH", "TL", "TM", "AE", "UZ", "VN", "YE"],
        "Africa": ["DZ", "AO", "BJ", "BW", "BF", "BI", "CM", "CV", "CF", "TD", "KM", "CG", "CD", "DJ", "EG", "GQ", "ER",
                   "ET", "GA", "GM", "GH", "GN", "GW", "CI", "KE", "LS", "LR", "LY", "MG", "MW", "ML", "MR", "MU", "MA",
                   "MZ", "NA", "NE", "NG", "RW", "ST", "SN", "SC", "SL", "SO", "ZA", "SS", "SD", "SZ", "TZ", "TG", "TN",
                   "UG", "ZM", "ZW"],
        "Europe": ["AL", "AD", "AT", "BY", "BE", "BA", "BG", "HR", "CY", "CZ", "DK", "EE", "FI", "FR", "DE", "GR", "HU",
                   "IS", "IE", "IT", "XK", "LV", "LI", "LT", "LU", "MT", "MC", "ME", "NL", "MK", "NO", "PL", "PT", "RO",
                   "RU", "SM", "RS", "SK", "SI", "ES", "SE", "CH", "TR", "UA", "GB", "VA"],
        "North America": ["AG", "BS", "BB", "BZ", "CA", "CR", "CU", "DM", "DO", "SV", "GD", "GT", "HT", "HN", "JM",
                          "MX", "NI", "PA", "KN", "LC", "VC", "TT", "US"],
        "South America": ["AR", "BO", "BR", "CL", "CO", "EC", "GY", "PY", "PE", "SR", "UY", "VE"],
        "Oceania": ["AU", "FJ", "KI", "MH", "FM", "NR", "NZ", "PW", "PG", "WS", "SB", "TO", "TV", "VU"],
        "Antarctica": ["AQ"]
    }
    
    # Convert the country name to a two-letter code
    if len(country_name) == 2:
        country_code = country_name
    elif len(country_name) == 3:
        country_code = pycountry.countries.get(alpha_3=country_name).alpha_2
    else:
        country_code = pycountry.countries.get(name=country_name).alpha_2
    
    # Find the region of the country
    for region, countries in continent_mapping.items():
        if country_code in countries:
            return region
    return None

def estimate_text_block_dimensions(text, font_size_pt):
    # Split text by '\n' to count lines
    lines = text.split("\n")
    num_lines = len(lines)
    max_line_length = max(len(line) for line in lines)
    
    # Font size in inches
    line_height = font_size_pt / 32  # Height per line
    char_width = line_height * 0.5  # Average character width
    
    # Total dimensions
    height = num_lines * line_height
    height = height * 1.02 if line_height <= 3 else height
    max_line_length = max_line_length * 0.9 if max_line_length > 20 else max_line_length
    max_line_length = max_line_length * 0.9 if max_line_length > 35 else max_line_length
    width = max_line_length * char_width
    return width * 1.02, height * 1.02

def estimate_max_dimensions(text_strs, font_size_pt):
    max_w, max_h = 0, 0
    for i in text_strs:
        w, h = estimate_text_block_dimensions(i, font_size_pt)
        if w > max_w:
            max_w = w
        if h > max_h:
            max_h = h
    return max_w, max_h

def generate_single_box(text, font_size_pt, if_save, save_path ,w = None, h = None):
    w0, h0 = estimate_text_block_dimensions(text, font_size_pt)
    if w is None:
        w = w0
    if h is None:
        h = h0
    with schemdraw.Drawing() as d:
        d.config(fontsize=font_size_pt)
        flow.Box(w=w, h=h).label(text, fontsize=font_size_pt, font='Arial')
    if if_save:
        d.save(save_path, dpi=300)
    return d

def generate_multiple_boxes(text_dict, font_size_pt, if_save, save_path, n_keep = 5):
    with schemdraw.Drawing() as d:
        d.config(fontsize= font_size_pt)
        # get max w and h of all keyword_dict
        max_w, max_h = 0, 0
        new_text_dict = {}
        for k in text_dict.keys():
            if len(text_dict.get(k)) > n_keep:
                temp_list = text_dict.get(k)[:n_keep] + ['...']
            else:
                temp_list = text_dict.get(k)
            new_text_dict[k] = '\n'.join(temp_list)
            w, h = estimate_text_block_dimensions('\n'.join(temp_list),  font_size_pt)
            if w > max_w:
                max_w = w
            if h > max_h:
                max_h = h
        count = 0
        for k in text_dict.keys():
            # Add boxes to c2 to the right, and ordered them from top to bottom
            w, h = estimate_text_block_dimensions('\n'.join(temp_list),  font_size_pt)
            c = flow.Box(w=max_w, h=max_h).at((0, (max_h + 0.2) * count)).label(new_text_dict.get(k), fontsize=12, font='Arial')
            # connect the box to c2
            d += c
            count += 1
    if if_save:
        d.save(save_path, dpi=300)    

def generate_strs(stat_val_dict, keyword_topics, keyword_subset, topic_selected):
    str1 = "Document retrieved \n from WOS query\n"
    topic_string = "Artificial Intelligence"
    Databases = ['SCI', 'SSCI']
    publication_num = stat_val_dict.get('TotalPublications')
    WOS_keywords_num = stat_val_dict.get('TotalKeywords')
    WOS_keywords_plus_num = stat_val_dict.get('TotalKeywordsPlus')
    single_author_num = stat_val_dict.get('SingleAuthoredDocuments')
    coauthor_mean = stat_val_dict.get('Co-authorsPerDocument')
    avg_citations = stat_val_dict.get('AverageCitations')
    avg_refs = stat_val_dict.get('AverageReferences')
    avg_document_age = stat_val_dict.get('AverageDocumentAge')
    total_country_num = stat_val_dict.get('TotalCountries')
    total_study_site_num = stat_val_dict.get('TotalStudySites')
    journal_num = stat_val_dict.get('Journals')
    growth_rate = stat_val_dict.get('AnnualGrowthRate')
    top_n = 10
    
    selected_topic_dict = {k: keyword_topics.get(k) for k in topic_selected}

    Text1_str = f"{str1}\nYear: {stat_val_dict.get('TimeSpan')}\nTopic: {topic_string}\nDatabases: {Databases}\nTotal Documents: {publication_num}"

    Topic_str = f"Topic Analysis\n" + '\n'.join([str(i) + f' ({len(keyword_subset.get(i))})' for i in selected_topic_dict.keys()])
    Topic_str += f"\n Other Topics ..."
    Keyword_str = f"Keyword Analysis\nKeyword ({WOS_keywords_num})\nKeyword Plus ({WOS_keywords_plus_num})"
    Sentiment_str = "Textblob Analysis\nTitle/Abstract Sentiment\n Title/Abstract Subjectivity"
    Paper_str = f"Paper Analysis\nSingle Author Num ({single_author_num})\nAverage Author Numbern ({coauthor_mean:.2f})\n\
        Average Citations ({avg_citations:.2f})\nAverage References ({avg_refs:.2f})\nAverage Document Age ({avg_document_age:.2f})\n\
            Publication Growth Rate ({growth_rate:.2f})"
    Country_str = f"Country Analysis\nAuthor Affiliated Country ({total_country_num})\nStudy Site ({total_study_site_num})"
    Journal_str = f"Journal Analysis\nJournal Number ({journal_num})\nTop {top_n} Journals"
    return [Text1_str, Topic_str, Keyword_str, Sentiment_str, Paper_str, Country_str, Journal_str], selected_topic_dict

def save_dict_into_docx_table(res_dict, doc_folder, doc_name):
    os.makedirs(doc_folder, exist_ok=True)
    doc = docx.Document()
    doc.add_heading('Meta-Analysis Results', level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for key, value in res_dict.items():
        row_cells = table.add_row().cells
        row_cells[0].text = key
        row_cells[1].text = str(value)
    doc.save(f'{doc_folder}{doc_name}.docx')

def write_supplement_doc(doc_folder, doc_name, supplement_doc_infos, topic_cats = None):
    os.makedirs(doc_folder, exist_ok=True)
    doc = docx.Document()
    
    run = doc.add_heading(level=0).add_run('Supplementary Tables')
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    for key in supplement_doc_infos:
        if key == 'Journal_Count':
            run = doc.add_heading(level=1).add_run('Top 10 Journal Total Count Information')
            run.font.color.rgb = RGBColor(0, 0, 0)
            table0 = doc.add_table(rows=0, cols=2)
            table0.style = 'Light Shading Accent 1'
            #jounal and count are two columns of pd
            temp_data = supplement_doc_infos[key]
            row_cells = table0.add_row().cells
            row_cells[0].text = 'Journal'
            row_cells[1].text = 'Count'
            for i in range(len(temp_data)):
                row_cells = table0.add_row().cells
                row_cells[0].text = str(temp_data.iloc[i, 0])
                row_cells[1].text = str(temp_data.iloc[i, 1])
            for row in table0.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
        elif key == 'Yearly_Journal_Count':
            run = doc.add_heading(level=1).add_run('Top 10 Journal Yearly Trends')
            run.font.color.rgb = RGBColor(0, 0, 0)
            temp_data = supplement_doc_infos[key]
            col_num = len(temp_data.columns) + 1
            table1 = doc.add_table(rows=0, cols=col_num)
            table1.style = 'Light Shading Accent 1'
            #jounal and count are two columns of pd
            
            row_cells = table1.add_row().cells
            row_cells[0].text = 'Year'
            for i in range(1, col_num):
                row_cells[i].text = temp_data.columns[i - 1]
            for j in range(len(temp_data)):
                row_cells = table1.add_row().cells
                row_cells[0].text = str(temp_data.index[j])
                for k in range(1, col_num):
                    # convert to int
                    row_cells[k].text = f"{temp_data.iloc[j, k - 1]:.0f}" 
            
            for row in table1.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
        elif key == 'Yearly_Publications':
            temp_data = supplement_doc_infos[key]
            run = doc.add_heading(level=1).add_run('Yearly Total and Countrywide Publication Count Information')
            run.font.color.rgb = RGBColor(0, 0, 0)
            years = list(temp_data.keys())
            years.sort()
            # only keep the ones with four digits
            years = [x for x in years if len(x) == 4]
            test_data = temp_data[years[0]]
            
            col_num = len(test_data.get('countries')) * 2 + 2
            table2 = doc.add_table(rows=0, cols=col_num)
            table2.style = 'Light Shading Accent 1'
            #jounal and count are two columns of pd
            countries = list(test_data.get('countries').keys())
            row_cells = table2.add_row().cells
            row_cells[0].text = 'Year'
            row_cells[1].text = 'Total'
            for j in range(2, col_num, 2):
                j0 = j // 2 - 1
                row_cells[j].text = countries[j0] + '-Count'
                row_cells[j + 1].text = countries[j0] + '-Percent'
            for year in years:
                row_cells = table2.add_row().cells
                row_cells[0].text = year
                test_data = temp_data[year]
                row_cells[1].text = str(test_data.get('total_publications'))
                for j in range(2, col_num, 2):
                    j0 = j // 2 - 1
                    row_cells[j].text = str(test_data.get('countries').get(countries[j0]).get('count'))
                    row_cells[j + 1].text = f"{test_data.get('countries').get(countries[j0]).get('percentage'):.2f}"
            for row in table2.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
        elif key == 'WOS_Keyword_trend':
            temp_data = supplement_doc_infos[key]
            run = doc.add_heading(level=1).add_run('Trend of Web of Science Keywords')
            run.font.color.rgb = RGBColor(0, 0, 0)
            col_num = len(temp_data) + 1
            table3 = doc.add_table(rows=0, cols=col_num)
            table3.style = 'Light Shading Accent 1'
            #jounal and count are two columns of pd
            row_cells = table3.add_row().cells
            row_cells[0].text = 'Year'
            temp_keys = list(temp_data.keys())
            years = list(temp_data.get(temp_keys[0]).keys())
            years.sort()
            # only keep the ones with four digits
            years = [x for x in years if x > 1000]
            for i in range(1, col_num):
                row_cells[i].text = temp_keys[i - 1].strip()
            for j in years:
                row_cells = table3.add_row().cells
                row_cells[0].text = str(j)
                for k in range(1, col_num):
                    row_cells[k].text = str(temp_data.get(temp_keys[k - 1]).get(j))
            for row in table3.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)    
        elif key == 'Country_sentiment' or key == 'Country_subjectivity':
            temp_data = supplement_doc_infos[key]
            # upper case the first letter
            key_suffix = key.split('_')[1].capitalize()
            run = doc.add_heading(level=1).add_run('Publications %s Information at Country Level' % key_suffix)
            run.font.color.rgb = RGBColor(0, 0, 0)
            col_num = 9
            table4 = doc.add_table(rows=0, cols=col_num)
            table4.style = 'Light Shading Accent 1'
            col_names = ['country', 'count', 'mean', 'std', 'min', '25%', '50%', '75%', 'max']
            row_cells = table4.add_row().cells
            for i in range(col_num):
                row_cells[i].text = col_names[i]
            for c in temp_data.keys():
                row_cells = table4.add_row().cells
                row_cells[0].text = c
                row_cells[1].text = f"{temp_data[c].get('count'):.0f}"
                row_cells[2].text = f"{temp_data[c].get('mean'):.2f}"
                row_cells[3].text = f"{temp_data[c].get('std'):.2f}"
                row_cells[4].text = f"{temp_data[c].get('min'):.2f}"
                row_cells[5].text = f"{temp_data[c].get('25%'):.2f}"
                row_cells[6].text = f"{temp_data[c].get('50%'):.2f}"
                row_cells[7].text = f"{temp_data[c].get('75%'):.2f}"
                row_cells[8].text = f"{temp_data[c].get('max'):.2f}"
                
            for row in table4.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)                    
        elif key == 'Country_Coop':
            run = doc.add_heading(level=1).add_run('Country Cooperation Information')
            run.font.color.rgb = RGBColor(0, 0, 0)
            temp_data = supplement_doc_infos[key]
            countries = list(temp_data.keys())
            # col_num needs to be set to the max number of sub_sub_keys:
            col_num = max([len(temp_data[c].get('Val').get(list(temp_data[c].get('Val'))[0])) for c in countries]) * 2 + 1
            # print(col_num)
            table5 = doc.add_table(rows=0, cols=col_num)
            table5.style = 'Light Shading Accent 1'
 
            for c in countries:
                temp_val_data = temp_data[c].get('Val')
                temp_percent_data = temp_data[c].get('Percent')
                row_cells = table5.add_row().cells
                row_cells[0].text = '\n' + c +'\n'
                
                sub_keys = list(temp_val_data.keys())
                sub_sub_keys = list(temp_val_data.get(sub_keys[0]).keys())
                # print(sub_sub_keys)
                row_cells = table5.add_row().cells
                row_cells[0].text = 'Year'
                for i in range(len(sub_sub_keys)):
                    # print(i)
                    row_cells[2 * i + 1].text = sub_sub_keys[i] + '-Val'
                    row_cells[2 * i + 2].text = sub_sub_keys[i] + '-Percent'
                    
                for k in temp_val_data.keys():
                    sub_temp_val_data = temp_val_data.get(k)
                    sub_temp_percent_data = temp_percent_data.get(k)
                    row_cells = table5.add_row().cells
                    row_cells[0].text = str(k)
                    for i in range(len(sub_sub_keys)):
                        sub_temp_single_val = sub_temp_val_data.get(sub_sub_keys[i])
                        sub_temp_single_percent = sub_temp_percent_data.get(sub_sub_keys[i])
                        if sub_temp_single_val is None:
                            sub_temp_single_val = 0
                        row_cells[2 * i + 1].text = str(sub_temp_single_val)
                        try:
                            row_cells[2 * i + 2].text = f"{sub_temp_single_percent:.2f}"
                        except:
                            row_cells[2 * i + 2].text = '0'
            
            for row in table5.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
            
            for p in range(len(countries)):
                a = table5.cell((p - 1) * (len(sub_keys)+2), 0)
                b = table5.cell((p - 1) * (len(sub_keys)+2), col_num - 1)
                a.merge(b)    
        elif key == 'Study Site Count' or key == 'Author Site Count':
            temp_data = supplement_doc_infos[key]
            run = doc.add_heading(level=1).add_run('Study Site Count by Country')
            run.font.color.rgb = RGBColor(0, 0, 0)
            # sort key based on the value
            temp_data = dict(sorted(temp_data.items(), key=lambda item: item[1], reverse=True))
            key_list = list(temp_data.keys())
            col_num = 8
            table6 = doc.add_table(rows=0, cols=col_num)
            table6.style = 'Light Shading Accent 1'
            row_cells = table6.add_row().cells
            for i in range(0, col_num, 2):
                row_cells[i].text = 'country'
                row_cells[i + 1].text = 'count'
            for i in range(0, len(key_list), 4):
                row_cells = table6.add_row().cells
                for j in range(4):
                    if i + j < len(key_list):
                        row_cells[j * 2].text = key_list[i + j]
                        row_cells[j * 2 + 1].text = str(temp_data[key_list[i + j]])  
            for row in table6.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
        else:
            temp_data = supplement_doc_infos[key]
            temp_trends = temp_data.get('Keyword_Trends')
            temp_cited_num = temp_data.get('Cited_Num')
            temp_sentiment = temp_data.get('Sentiment')
            temp_subjectivity = temp_data.get('Subjectivity')
            # upper case the first letter
            if topic_cats is not None:
                key_suffix = topic_cats.get(key)
            else:
                key_suffix = key.capitalize()
            if not temp_trends:
                continue 
            # first, deal with keyword trends
            run = doc.add_heading(level=1).add_run('The Publication Trends for %s' % key_suffix)
            run.font.color.rgb = RGBColor(0, 0, 0)
            col_num = len(temp_trends) + 1
            table8 = doc.add_table(rows=0, cols=col_num)
            table8.style = 'Light Shading Accent 1'
            #jounal and count are two columns of pd
            row_cells = table8.add_row().cells
            row_cells[0].text = 'Year'
            temp_keys = list(temp_trends.keys())
            years = list(temp_trends.get(temp_keys[0]).keys())
            years.sort()
            years = [x for x in years if x > 1000]
            for i in range(1, col_num):
                row_cells[i].text = temp_keys[i - 1].replace(' ', '')
            for j in years:
                row_cells = table8.add_row().cells
                row_cells[0].text = str(j)
                for k in range(1, col_num):
                    row_cells[k].text = str(temp_trends.get(temp_keys[k - 1]).get(j))
            for row in table8.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8)  
            
            # second, deal with the remaining
            col_num = 9
            row_num = len(temp_cited_num)
            temp_keys = list(temp_cited_num.keys())
            
            run = doc.add_heading(level=1).add_run('The Publication Citation, Sentiment, and Subjectivity for %s' % key_suffix)
            run.font.color.rgb = RGBColor(0, 0, 0)
            
            table9 = doc.add_table(rows=0, cols=col_num)
            table9.style = 'Light Shading Accent 1'
            
            row_cells = table9.add_row().cells
            row_cells[0].text = 'Topic'
            row_cells[1].text = 'Count'
            row_cells[2].text = 'Citation'
            row_cells[3].text = 'Sentiment_Mean'
            row_cells[4].text = 'Sentiment_Min'
            row_cells[5].text = 'Sentiment_Max'
            row_cells[6].text = 'Subjectivity_Mean'
            row_cells[7].text = 'Subjectivity_Min'
            row_cells[8].text = 'Subjectivity_Max'
            
            for i in range(row_num):
                row_cells = table9.add_row().cells
                temp_key = temp_keys[i]
                row_cells[0].text = temp_key
                row_cells[1].text = f"{temp_sentiment.get(temp_key).get('count'):.0f}"
                row_cells[2].text = f"{temp_cited_num.get(temp_key):.2f}"
                row_cells[3].text = f"{temp_sentiment.get(temp_key).get('mean'):.2f}"
                row_cells[4].text = f"{temp_sentiment.get(temp_key).get('min'):.2f}"
                row_cells[5].text = f"{temp_sentiment.get(temp_key).get('max'):.2f}"
                row_cells[6].text = f"{temp_subjectivity.get(temp_key).get('mean'):.2f}"
                row_cells[7].text = f"{temp_subjectivity.get(temp_key).get('min'):.2f}"
                row_cells[8].text = f"{temp_subjectivity.get(temp_key).get('max'):.2f}"
                
            for row in table9.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(8) 
    # set margin of doc to narrow
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)
    doc.save(f'{doc_folder}{doc_name}.docx')

def standardize_column(column):
    # Create a mapping of normalized values (ignoring spaces) to the longest version
    mapping = {}
    for value in column.dropna():  # Exclude missing values (NaN)
        normalized = value.replace(" ", "")  # Normalize by removing all spaces
        if normalized not in mapping or len(value.split()) > len(mapping[normalized].split()):
            mapping[normalized] = value  # Keep the one with the most words
    
    # Replace values in the column using the mapping, keeping NaN intact
    return column.apply(lambda x: mapping[x.replace(" ", "")] if pd.notna(x) else x)

def estimate_growth_rate(x, y):
    y0 = np.log(np.array(y)) - np.log(y[0])
    x0 = np.array(x)
    slope, intercept, r_value, p_value, std_err = linregress(x0, y0)
    growth_rate = np.exp(slope) - 1
    return growth_rate

def generate_stats(combined_df, start_year, end_year, yearly_report,\
    author_country_count_dict, study_site_count_dict):
    total_publications = len(combined_df)
    total_country_num = len(author_country_count_dict)
    study_site_count_num = len(study_site_count_dict)
    journals = combined_df['journal']
    # make every elements in lower case
    journals = journals.apply(lambda x: x.lower() if isinstance(x, str) else 'NaN')
    journal_unique_num = len(set(journals)) - 1
    total_years = end_year - start_year + 1
    x_reg, y_reg = [], []
    for i in range(start_year, end_year + 1):
        x_reg.append(i - start_year)
        y_reg.append(yearly_report.get(str(i)).get('total_publications'))
    average_growth_rate = 100*estimate_growth_rate(x_reg, y_reg)
    # for i in range(start_year, end_year - 1):
    #     growth_rate = (yearly_report.get(str(i+1)).get('total_publications') - yearly_report.get(str(i)).get('total_publications'))/total_years
    #     growth_rates.append(growth_rate)
    # average_growth_rate = sum(growth_rates)/len(growth_rates)
    total_citations = int(combined_df['times-cited'].sum())
    total_refs = int(combined_df['number-of-cited-references'].sum())
    average_citations = total_citations/total_publications
    average_refs = total_refs/total_publications
    average_document_age = combined_df['year'].apply(lambda x: end_year - int(x)).mean()
    
    # go to country list, to see if the country is a list of one
    single_author_num = 0
    total_author_num = 0
    coop_author_num = 0
    coop_authors = []
    # search the country column of combined_df
    for i in combined_df['country']:
        temp_list = ast.literal_eval(i)
        if len(temp_list) == 0:
            continue
        else:
            total_author_num += 1
            if len(temp_list) == 1:
                single_author_num += 1
            else:
                coop_authors.append(len(temp_list))
                # print(list(set(temp_list)))
                if len(list(set(temp_list))) > 1:
                    coop_author_num += 1
    # single_author_ratio = single_author_num/total_author_num
    # print(coop_author_num, total_author_num)
    coop_author_ratio = coop_author_num/total_author_num * 100
    coop_author_mean = sum(coop_authors) / len(coop_authors) 
    keyword_list, keyword_plus_list = [], []
    for i in combined_df['keywords']:
        if isinstance(i, str):
            keyword_list.extend(i.split(';'))
    
    for i in combined_df['keywords-plus']:
        if isinstance(i, str):
            keyword_plus_list.extend(i.split(';'))
    # print(keyword_list)
    keyword_num = len(list(set(keyword_list)))
    keyword_plus_num = len(list(set(keyword_plus_list)))
    res_val_dict = {
        'TimeSpan': f'{start_year} - {end_year}',
        'TotalPublications': total_publications,
        'Journals': journal_unique_num,
        'AnnualGrowthRate': average_growth_rate,
        'AverageDocumentAge': average_document_age,
        'TotalCitations': total_citations,
        'AverageCitations': average_citations,
        'TotalReferences': total_refs,
        'AverageReferences': average_refs,
        'TotalCountries': total_country_num,
        'TotalStudySites': study_site_count_num,
        'SingleAuthoredDocuments': single_author_num,
        'Co-authorsPerDocument': coop_author_mean,
        'InternationalCollaborationRatio': coop_author_ratio,
        'TotalKeywords': keyword_num,
        'TotalKeywordsPlus': keyword_plus_num}
    res_dict = {
        'Description': 'Results',
        '  Time Span': f'{start_year} - {end_year}',
        '  Total Publications': f"{total_publications:d}",
        '  Journals': f"{journal_unique_num:d}",
        '  Annual Growth Rate (%)': f"{average_growth_rate:.2f}",
        '  Average Document Age (yr)': f"{average_document_age:.2f}",
        'Citation and Reference': '',
        '  Total Citations': f"{total_citations:d}",
        '  Average Citations': f"{average_citations:.2f}",
        '  Total References': f"{total_refs:d}",
        '  Average References': f"{average_refs:.2f}",
        'Geographical Distribution': '',
        '  Total Countries': f"{total_country_num:d}",
        '  Total Study Sites': f"{study_site_count_num:d}",
        'Authors': '',
        '  Single-authored Documents': f"{single_author_num:d}",
        '  Co-authors Per Document': f"{coop_author_mean:.2f}",
        '  International Collaboration Ratio': f"{coop_author_ratio:.2f}%",
        'Keywords': '',
        '  Total Keywords': f"{keyword_num:d}",
        '  Total Keywords Plus': f"{keyword_plus_num:d}"
        
    }
    return res_dict, res_val_dict

def train_word_vec_model(combined_df):
    corpus = list(combined_df.abstract.dropna().values)
    # Tokenize the sentences into words
    print("Tokenizing corpus...")
    tokenized_corpus = [sentence.lower().split() for sentence in tqdm(corpus)]

    # Count word frequencies
    print("Counting word frequencies...")
    word_counts = Counter(word for sentence in tokenized_corpus for word in sentence)

    # Train Word2Vec model
    print("Training Word2Vec model...")
    model = Word2Vec(sentences=tokenized_corpus, vector_size=50, window=3, min_count=1, workers=4)
    
    return model, word_counts

def create_KNN_PCA_res(word_vectors, keyword_list, n_components, num_clusters, random_state = 45):
    

    # Filter word vectors to include only words in the keyword list
    print("Filtering word vectors based on the provided keyword list...")
    filtered_words = [word for word in tqdm(keyword_list) if word in word_vectors]
    filtered_vectors = [word_vectors[word] for word in filtered_words]
    # Perform K-means clustering
    print(f"Clustering into {num_clusters} clusters...")
    kmeans = KMeans(n_clusters=num_clusters, random_state=random_state)
    #kmeans = KMeans(n_clusters=num_clusters, random_state=42)
    kmeans.fit(filtered_vectors)

    # Map words to clusters
    word_clusters = {filtered_words[i]: kmeans.labels_[i] for i in range(len(filtered_words))}

    # Reduce dimensions for visualization
    print(f"Reducing dimensions to {n_components} using PCA...")
    pca = PCA(n_components=n_components)
    reduced_vectors = pca.fit_transform(filtered_vectors)
    return reduced_vectors, word_clusters, kmeans, pca

def update_keyword_subset(
    df,
    keyword_subset,
    title_column="title",
    abstract_column="abstract",
    word_1=None,
    word_2=None,
    opr="AND",
    key_to_add=None
):
    """
    Updates the keyword subset based on the specified word matching logic.

    :param df: DataFrame containing the data.
    :param keyword_subset: Existing keyword subset dictionary.
    :param title_column: Column name for titles (default: "title").
    :param abstract_column: Column name for abstracts (default: "abstract").
    :param word_1: The primary word to match.
    :param word_2: The secondary word for additional logic.
    :param opr: Operation between word_1 and word_2 ('AND' or 'NOT').
    :param key_to_add: Key in the keyword subset to add the new indices.
    :return: Updated keyword subset.
    """
    if not word_1 or not key_to_add:
        raise ValueError("Both 'word_1' and 'key_to_add' must be provided.")

    word_1 = word_1.lower()
    word_2 = word_2.lower() if word_2 else None

    # Initialize a set to collect new indices
    new_indices = set()

    for idx, row in tqdm(df.iterrows(), desc = 'Process Rows', total = len(df)):
        # Combine title and abstract for search
        text = f"{row[title_column]} {row[abstract_column]}".lower()

        # Check for matches based on the operation
        if opr == "AND" and word_2:
            # Both word_1 and word_2 must be present
            if re.search(rf'\b{re.escape(word_1)}\b', text) and re.search(rf'\b{re.escape(word_2)}\b', text):
                new_indices.add(idx)
        elif opr == "NOT" and word_2:
            # word_1 must exist, but word_2 must not be in association with word_1
            if re.search(rf'\b{re.escape(word_1)}\b', text):
                match_word_1 = re.finditer(rf'\b{re.escape(word_1)}\b', text)
                match_word_2 = re.finditer(rf'\b{re.escape(word_2)}\b', text)

                # Check if any match of word_1 overlaps with word_2
                word_2_positions = {m.start() for m in match_word_2}
                if not any(m.start() in word_2_positions for m in match_word_1):
                    new_indices.add(idx)
        elif opr == "AND" and not word_2:  # Single-word case, equivalent to just matching word_1
            if re.search(rf'\b{re.escape(word_1)}\b', text):
                new_indices.add(idx)
        else:
            raise ValueError("Invalid operation. Supported operations are 'AND' and 'NOT'.")

    # Merge new indices with the existing key in the keyword subset
    if key_to_add in keyword_subset:
        keyword_subset[key_to_add] = list(set(keyword_subset[key_to_add]) | new_indices)
    else:
        keyword_subset[key_to_add] = list(new_indices)

    return keyword_subset

def get_major_coop_trends(ranged_country_pairs):
    EU = ["AL", "AD", "AT", "BY", "BE", "BA", "BG", "HR", "CY", "CZ", "DK", "EE", "FI", "FR", "DE", "GR", "HU",\
                   "IS", "IE", "IT", "XK", "LV", "LI", "LT", "LU", "MT", "MC", "ME", "NL", "MK", "NO", "PL", "PT", "RO",\
                   "SM", "RS", "SK", "SI", "ES", "SE", "CH", "TR", "UA", "VA"]
    LAC = ["AR", "BO", "BR", "CL", "CO", "CR", "CU", "DO", "EC", "MX", "SV", "GT", "HT", "PA", "HN", "MX", "NI", "PY", "PE", "PR", "UY", "VE"]
    RM = ["RU", "GB", "US", "CN", "IN", "JP", "KR", "SG", "CA"]
    OC = ["AU", "NZ"]
    res = {}
    for k in ranged_country_pairs.keys():
        temp_res = {}
        temp_dict = ranged_country_pairs.get(k)
        for k0 in temp_dict.keys():
            country_1, country_2 = k0[0], k0[1]
            temp_val = temp_dict.get(k0)
            if (country_1 in EU and country_2 in EU) or (country_1 in OC and country_2) in OC or (country_1 in LAC and country_2 in LAC):
                continue
            elif country_1 not in EU and country_1 not in OC and country_1 not in LAC and country_1 not in RM:
                continue
            elif country_2 not in EU and country_2 not in OC and country_2 not in LAC and country_2 not in RM:
                continue
            elif country_1 in EU:
                if country_2 in RM:
                    temp_key = (country_2, 'EU') if country_2 == 'CN' else ('EU', country_2)
                elif country_2 in OC:
                    temp_key = ('EU', 'OC')
                elif country_2 in LAC:
                    temp_key = ('EU', 'LAC')
            elif country_2 in EU:
                if country_1 in RM:
                    temp_key = (country_1, 'EU') if country_1 == 'CN' else ('EU', country_1)
                elif country_1 in OC:
                    temp_key = ('EU', 'OC')
                elif country_1 in LAC:
                    temp_key = ('EU', 'LAC')
                    
            elif country_1 in OC:
                if country_2 in RM:
                    temp_key = ('OC', country_2) if country_2 == 'RU' or country_2 == 'US' or country_2 == 'SG' else (country_2, 'OC')
                elif country_2 in LAC:
                    temp_key = ('LAC', 'OC')
            elif country_2 in OC:
                if country_1 in RM:
                    temp_key = ('OC', country_1) if country_1 == 'RU' or country_2 == 'US' or country_2 == 'SG' else (country_1, 'OC')
                elif country_1 in LAC:
                    temp_key = ('LAC', 'OC')
            elif country_1 in LAC:
                if country_2 in RM:
                    temp_key = ('LAC', country_2) if country_2 == 'RU' or country_2 == 'US' or country_2 == 'SG' else (country_2, 'LAC')
            elif country_2 in LAC:
                if country_1 in RM:
                    temp_key = ('LAC', country_1) if country_1 == 'RU' or country_1 == 'US' or country_1 == 'SG' else (country_1, 'LAC')   
            elif country_1 in RM and country_2 in RM:
                temp_key = (country_1, country_2) if country_1 < country_2 else (country_2, country_1)
            else:
                continue
            if temp_key not in temp_res:
                temp_res[temp_key] = temp_val
            else:
                temp_res[temp_key] += temp_val
        res[k] = temp_res
    return res

def get_single_country_coop_res(ranged_country_pairs, searched_country, is_percent = True):
    res = {}
    
    for k in ranged_country_pairs.keys():
        temp_res = {}
        temp_dict = ranged_country_pairs.get(k)
        for k0 in temp_dict.keys():
            country_1, country_2 = k0[0], k0[1]
            if country_1 == searched_country:
                paired_country = country_2
            elif country_2 == searched_country:
                paired_country = country_1
            else:
                continue
            temp_val = temp_dict.get(k0)
            if paired_country not in temp_res:
                temp_res[paired_country] = temp_val
            else:
                temp_res[paired_country] += temp_val
        if is_percent:
            total_val = sum(temp_res.values())
            temp_res = {k1: 100*v1/total_val for k1, v1 in temp_res.items()}
        res[k] = temp_res
    return res

def analyze_country_pairs_by_year_range(df, year_ranges, country_column):
    """
    Analyzes a DataFrame to compute country pair counts for given year ranges.

    :param df: Combined DataFrame containing year and country data
    :param year_ranges: List of tuples representing year ranges [(s1, e1), (s2, e2), ...]
    :param country_column: Column name containing country lists
    :return: Dictionary with year ranges as keys and country pair counts as values
    """
    range_country_pairs = {}

    # Ensure the year column is converted to integer
    
    for start_year, end_year in year_ranges:
        # Subset the DataFrame for the current year range
        subset_df = filter_df_based_on_year_range(df, 'year', start_year, end_year)

        # Compute country pair counts for the subset
        country_pair_dict = analyze_country_pairs(subset_df, country_column)

        # Add the result to the dictionary
        range_country_pairs[(start_year, end_year)] = country_pair_dict

    return range_country_pairs

def filter_keyword_subset_by_year(df, keyword_subset, year_column, start_year, end_year):
    """
    Filters the keyword subset to include only entries within a specified year range.

    :param df: DataFrame containing the year information.
    :param keyword_subset: Dictionary where keywords map to lists of row indices.
    :param year_column: Column name in the DataFrame containing year data.
    :param start_year: Start year as a string (format: 'yyyy').
    :param end_year: End year as a string (format: 'yyyy').
    :return: A filtered keyword subset dictionary.
    """

    # Filter the DataFrame to get the indices in the year range
    valid_indices = set(df[(df[year_column] >= start_year) & (df[year_column] <= end_year)].index)

    # Filter the keyword subset
    filtered_keyword_subset = {
        keyword: [idx for idx in indices if idx in valid_indices]
        for keyword, indices in keyword_subset.items()
    }

    return filtered_keyword_subset

def get_subset_of_country_pairs(combined_df, study_site_subset, country_1, country_2, keyword_dict):
    subsets_to_check = get_subset_of_keys(study_site_subset, [country_1, country_2])
    # get the rows occurred simultaneously in both countries
    rows1, rows2 = subsets_to_check[country_1], subsets_to_check[country_2]
    rows_to_check = list(set(rows1).intersection(rows2))
    # combined_df filter by rows_to_check
    subset_df = combined_df.loc[rows_to_check]
    # get the keyword subset
    keyword_subset = create_keyword_subset(subset_df, keyword_dict, title_column="title", abstract_column="abstract")
    return keyword_subset

def analysis_multi_val_column(df, column_name, separator=';'):
    res_dict = {}
    for i, row in tqdm(df.iterrows(), total=df.shape[0], desc="Processing rows"):
        if not isinstance(row[column_name], str):
            continue
        values = row[column_name].split(separator)
        for value in values:
            if not res_dict or value not in res_dict:
                res_dict[value] = [i]
            else:
                res_dict[value].append(i)
    return res_dict

def keep_top_n_subsets(subset_dict, n=5):
    res_dict = {}
    # sort the subsets by size
    sorted_subsets = sorted(subset_dict.items(), key=lambda x: len(x[1]), reverse=True)
    # keep the top n subsets
    for i in range(min(n, len(sorted_subsets))):
        res_dict[sorted_subsets[i][0]] = sorted_subsets[i][1]
    # make the word for key, only the first letter is upper case
    res_dict = {k.capitalize(): v for k, v in res_dict.items()}
    return res_dict

def extract_names(input_string):
    """
    Extracts names from a string with the format 'Name/ID'.

    Parameters:
        input_string (str): The input string containing names and IDs.

    Returns:
        list: A list of extracted names.
    """
    # Regular expression to match names (everything before the '/')
    pattern = r'(.*?)/[A-Za-z0-9-]+'
    matches = re.findall(pattern, input_string)
    return [name.strip() for name in matches]

def extract_subset_values(df, keyword_subset, column_name):
    subset_data = []
    for keyword, indices in keyword_subset.items():
        subset_df = df.loc[indices]
        #print(len(subset_df))
        if not subset_df.empty:
            for val in subset_df[column_name]:
                subset_data.append({"Keyword": keyword, "Value": val})
    subset_df = pd.DataFrame(subset_data)
    means = subset_df.groupby("Keyword")["Value"].mean()
    return means

def create_gender_subsets(combined_df, name_dict, name_column='orcid-numbers'):
    res_dict = {'boy':[], 'girl':[]}
    for i, row in tqdm(combined_df.iterrows(), total=combined_df.shape[0], desc="Processing rows"):
        if not isinstance(row[name_column], str):
            continue
        names = extract_names(row[name_column])
        is_boy, is_girl = False, False
        if names:
            for name in names:
                name_to_check = name.split(',')[-1].strip()
                if name_to_check in name_dict:
                    gender = name_dict[name_to_check]
                    if gender == 'Boy':
                        is_boy = True
                    elif gender == 'Girl':
                        is_girl = True
        if is_boy:
            res_dict['boy'].append(i)
        if is_girl:
            res_dict['girl'].append(i)
    return res_dict
            
def get_doi_values(df, columns_to_check, words_to_check):
    """
    Returns a list of values from the 'doi' column if all words in `words_to_check` are 
    found across the specified columns in a row, after dropping rows with missing values
    in the columns to check.

    Parameters:
    - df (pd.DataFrame): The input DataFrame.
    - columns_to_check (list): List of column names to check for words.
    - words_to_check (list): List of words to search for.

    Returns:
    - list: List of 'doi' column values where the condition is met.
    """
    # Drop rows with missing values in the columns to check
    df = df.dropna(subset=columns_to_check)

    # Define a helper function to check if all words are in the row's combined columns
    def row_contains_all_words(row):
        combined_text = ' '.join(str(row[col]) for col in columns_to_check if col in df)
        return all(word in combined_text for word in words_to_check)
    
    # Filter the DataFrame for rows that meet the condition
    filtered_df = df[df.apply(row_contains_all_words, axis=1)]
    
    res_dict = {}
    # use unique-id as key, title as one value, doi as another value
    for i, row in filtered_df.iterrows():
        res_dict[row['unique-id']] = {'title':row['title'], 'doi':row['doi']}
    # Return the 'doi' column as a list
    return res_dict

def test_median_differences(combined_df, keyword_subset, column_name):
    group_names = list(keyword_subset.keys())
    group_data = []

    print(f"\n=== Group Summary for '{column_name}' ===")
    for key in group_names:
        values = combined_df.loc[keyword_subset[key], column_name].dropna().values
        group_data.append(values)
        mean_val = np.mean(values)
        median_val = np.median(values)
        print(f"{key}: mean = {mean_val:.4f}, median = {median_val:.4f}, n = {len(values)}")

    # Kruskal–Wallis test
    h_stat, p_kruskal = kruskal(*group_data)
    print(f"\nKruskal–Wallis test: H = {h_stat:.4f}, p = {p_kruskal:.4f}")

    if p_kruskal < 0.05:
        print("=> Significant difference detected among groups. Running pairwise Mann–Whitney U tests (Bonferroni corrected):")
        pairs = list(combinations(range(len(group_data)), 2))
        p_vals = []
        for i, j in pairs:
            _, p = mannwhitneyu(group_data[i], group_data[j], alternative='two-sided')
            p_vals.append(p)
        
        # Adjust p-values with Bonferroni correction
        _, p_adj, _, _ = multipletests(p_vals, method='bonferroni')
        
        for (i, j), p in zip(pairs, p_adj):
            sig_mark = " (significant)" if p < 0.05 else ""
            print(f"  {group_names[i]} vs {group_names[j]}: p_adj = {p:.4f}{sig_mark}")
    else:
        print("=> No significant difference in medians among groups.")
def test_group_differences(combined_df, keyword_subset, column_name):
    # Collect group data
    groups = []
    for key, indices in keyword_subset.items():
        group_values = combined_df.loc[indices, column_name].dropna().values
        groups.append(group_values)

    # ANOVA (parametric)
    f_stat, p_anova = f_oneway(*groups)

    # Kruskal-Wallis (non-parametric)
    h_stat, p_kruskal = kruskal(*groups)

    return {
        "ANOVA": {"F-statistic": f_stat, "p-value": p_anova},
        "Kruskal-Wallis": {"H-statistic": h_stat, "p-value": p_kruskal}
    }